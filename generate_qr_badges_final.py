#!/usr/bin/env python3
# v3.6.1-postfix — logos loaded from PNG files

"""
generate_qr_name_badges.py v3.6.1-postfix

This version loads the two PNG assets (full GT logo and small GT mark)
from external files at runtime. Place `GT_full_logo.png` and
`GT_small_logo.png` alongside this script. They will be opened with PIL
and displayed in the user interface.
"""

import sys
import os
import subprocess
import tempfile
import logging
import webbrowser
import argparse
import traceback
import time
import base64
from io import BytesIO
from pathlib import Path

# GUI imports
try:
    import tkinter as tk
    from tkinter import Tk, filedialog, messagebox, StringVar, font as tkfont
    from tkinter.ttk import Progressbar, Style, Button as TtkButton
    _tk_available = True
except ImportError:
    _tk_available = False

# PIL import for image handling
try:
    from PIL import Image, ImageTk
    # Pillow 10+ moved resampling filters under Image.Resampling
    if hasattr(Image, "Resampling"):
        RESAMPLE = Image.Resampling.LANCZOS
    else:  # pragma: no cover - older Pillow versions
        RESAMPLE = getattr(Image, "LANCZOS", Image.ANTIALIAS)
    _pil_available = True
except ImportError:
    _pil_available = False

# python-docx imports
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.enum.table import WD_ROW_HEIGHT_RULE
    from docx.enum.section import WD_ORIENT
except ImportError:
    print("Error: The 'python-docx' package is required. Install with:\n    pip install python-docx")
    sys.exit(1)

# pandas is required
try:
    import pandas as pd
except ImportError:
    print("Error: The 'pandas' package is required. Install with:\n    pip install pandas")
    sys.exit(1)

# Configure logging
default_format = '[%(asctime)s] [%(levelname)s] %(message)s'
logging.basicConfig(level=logging.INFO, format=default_format, datefmt='%Y-%m-%d %H:%M:%S')

# ----------------------------------------------------------
# 1) Base64‐encoded placeholders (replace with actual data)
# ----------------------------------------------------------

# Replace the string below with the actual Base64 of "GT_full_logo.png"
FULL_LOGO_B64 = ""
SMALL_LOGO_B64 = ""
RIBBON_B64 = ""
ASSET_DIR = Path(__file__).resolve().parent
FULL_LOGO_FILE = ASSET_DIR / "GT_full_logo.png"
SMALL_LOGO_FILE = ASSET_DIR / "GT_small_logo.png"
RIBBON_FILE = ASSET_DIR / "GT_ribbon.png"
def _load_embedded_image(b64_data):
    """
    Decode Base64 string into a PIL Image. Returns a PIL Image or None on failure.
    """
    if not _pil_available or not b64_data:
        return None
    try:
        img_bytes = base64.b64decode(b64_data)
        return Image.open(BytesIO(img_bytes))
    except Exception:
        return None


def _load_image_file(path: Path):
    """Load an image from the given file path. Returns a PIL Image or None."""
    if not _pil_available:
        return None
    try:
        return Image.open(path)
    except Exception:
        return None

# Georgia Tech brand colors
GT_GOLD = '#B3A369'
GT_NAVY = '#003057'
WHITE = '#FFFFFF'
GT_HOVER = '#D4C58C'  # lighter gold for hover

# Preferred font stack
PREFERRED_FONT = 'Roboto'
FALLBACK_FONTS = ['Helvetica', 'Arial']

DEFAULT_CONFIG = {
    'margins': {'top': 0.5, 'bottom': 0.5, 'left': 0.25, 'right': 0.25},
    'per_page_qr': 8,
    'columns_qr': 2,
    'qr_size_in': 1.5,
    'font_size_pt': 12,
    'label_dims_in_qr': {'width': 3.5, 'height': 2.5},
    'per_page_name': 6,
    'columns_name': 2,
    'label_dims_in_name': {'width': 4.0, 'height': 3.0}
}

REQUIRED_CSV_COLUMNS = [
    'Preferred', 'Last',
    'FASET Total Guest Count',
    'Guest 1 Preferred Name', 'Guest 1 Last Name', 'Guest 1 Affiliations',
    'Guest 2 Preferred Name', 'Guest 2 Last Name', 'Guest 2 Affiliations'
]


def _choose_font(root):
    available = set(tkfont.families(root))
    if PREFERRED_FONT in available:
        return (PREFERRED_FONT, 12)
    for fb in FALLBACK_FONTS:
        if fb in available:
            return (fb, 12)
    return ('TkDefaultFont', 12)


def parse_args():
    parser = argparse.ArgumentParser(description='Generate QR and various Name badges from CSV')
    parser.add_argument('-i', '--input', help='Input CSV file', metavar='PATH')
    parser.add_argument('-o', '--output', help='Output Word file (DOCX)', metavar='PATH')
    parser.add_argument('-c', '--config', help='Config file (JSON or YAML)', metavar='PATH')
    parser.add_argument('--landscape', action='store_true', help='Use landscape layout')
    return parser.parse_args()


def load_config(path: Path) -> dict:
    """Load configuration from JSON or YAML file."""
    if not path:
        return {}
    try:
        with open(path, 'r', encoding='utf-8') as fh:
            if path.suffix.lower() in {'.yaml', '.yml'}:
                try:
                    import yaml
                except ImportError:
                    logging.error('pyyaml is required for YAML configs.')
                    return {}
                return yaml.safe_load(fh) or {}
            else:
                import json
                return json.load(fh)
    except FileNotFoundError:
        logging.error('Config file not found: %s', path)
    except Exception:
        logging.exception('Failed to load config')
    return {}


def open_file(path: Path, background: bool = False) -> None:
    try:
        if sys.platform == 'darwin':
            cmd = ['open'] + (['-g'] if background else []) + [str(path)]
            subprocess.run(cmd, check=False)
        elif sys.platform == 'win32':
            os.startfile(path)
        else:
            webbrowser.open(path.as_uri())
    except Exception:
        logging.exception(f"Could not open {path}")


def pick_input_file() -> Path:
    if not _tk_available:
        print('Error: tkinter is not available. Supply -i <input.csv>.')
        sys.exit(1)
    root = Tk()
    root.withdraw()
    path = filedialog.askopenfilename(
        title='Select a CSV file',
        filetypes=[('CSV Files', '*.csv')]
    )
    root.destroy()
    if not path:
        sys.exit(0)
    return Path(path)


def pick_save_file(default_name: str) -> Path:
    if not _tk_available:
        print('Error: tkinter not available. Supply -o <output.docx>.')
        sys.exit(1)
    root = Tk()
    root.withdraw()
    save = filedialog.asksaveasfilename(
        title='Save Word file as…',
        defaultextension='.docx',
        initialfile=default_name,
        filetypes=[('Word Documents', '*.docx'), ('All Files', '*.*')]
    )
    root.destroy()
    if not save:
        sys.exit(0)
    return Path(save)


def load_records(file_path: Path):
    if file_path.suffix.lower() != '.csv':
        if _tk_available:
            messagebox.showerror('Invalid Input', 'Only CSV is supported.')
        else:
            print('Error: Only CSV input is supported.')
        sys.exit(1)

    try:
        df = pd.read_csv(file_path, dtype=str, encoding='utf-8-sig').fillna('')
    except Exception:
        logging.exception('Failed to read CSV.')
        sys.exit(1)

    missing = [col for col in REQUIRED_CSV_COLUMNS if col not in df.columns]
    if missing:
        msg = f"Missing required CSV columns: {', '.join(missing)}"
        if _tk_available:
            messagebox.showerror('Invalid CSV Schema', msg)
        else:
            print(f'Error: {msg}')
        sys.exit(1)

    return df.to_dict(orient='records')


def safe_int(val_str):
    """
    Try converting to int, default to 0 if empty or non-numeric.
    """
    try:
        return int(val_str)
    except ValueError:
        return 0


def download_qr_image(url: str, idx: int, timeout: int = 15):
    try:
        import requests
    except ImportError:
        logging.error("The 'requests' package is required for QR download. Skipping QR images.")
        return None
    try:
        resp = requests.get(url, timeout=timeout)
        resp.raise_for_status()
        tmp = tempfile.NamedTemporaryFile(suffix='.png', delete=False)
        tmp.write(resp.content)
        tmp.close()
        return Path(tmp.name)
    except Exception:
        logging.exception(f"QR #{idx} download failed")
        return None


def convert_to_pdf(docx_path: Path, outdir: Path) -> Path:
    pdf_path = outdir / (docx_path.stem + '.pdf')
    try:
        if sys.platform in ['darwin', 'win32']:
            from docx2pdf import convert
            convert(str(docx_path), str(pdf_path))
        else:
            subprocess.run([
                'libreoffice', '--headless', '--convert-to', 'pdf',
                str(docx_path), '--outdir', str(outdir)
            ], check=True)
        return pdf_path
    except ImportError:
        logging.error("The 'docx2pdf' package is required on Windows/macOS for PDF. Skipping PDF conversion.")
        return None
    except Exception:
        logging.exception('PDF conversion failed')
        return None


def create_progress_window(title: str, total: int):
    """Create and return a progress window with a bar and labels."""
    if not _tk_available:
        return None, None, None, None

    prog = Tk()
    prog.title(title)
    prog.configure(bg=GT_NAVY)
    prog.resizable(False, False)
    prog.attributes('-topmost', True)

    img_small = _load_image_file(SMALL_LOGO_FILE) or _load_embedded_image(SMALL_LOGO_B64)
    if img_small:
        logo_resized = img_small.resize((100, 50), RESAMPLE)
        logo_tk = ImageTk.PhotoImage(logo_resized)
        logo_lbl = tk.Label(prog, image=logo_tk, bg=GT_NAVY)
        logo_lbl.image = logo_tk
        logo_lbl.pack(pady=(10, 5))
    else:
        fallback_lbl = tk.Label(
            prog, text='GT Logo', bg=GT_NAVY, fg=WHITE, font=_choose_font(prog)
        )
        fallback_lbl.pack(pady=(10, 5))

    style = Style(prog)
    style.theme_use('clam')
    style.configure(
        'Flat.Horizontal.TProgressbar',
        troughcolor=GT_NAVY,
        background=GT_GOLD,
        thickness=20,
        bordercolor=GT_NAVY,
    )

    bar = Progressbar(
        prog,
        orient='horizontal',
        length=360,
        mode='determinate',
        maximum=total,
        style='Flat.Horizontal.TProgressbar',
    )
    bar.pack(padx=20, pady=(10, 5))

    bold_font = (*_choose_font(prog), 'bold')
    percent_label = tk.Label(
        prog,
        text='0%',
        bg=GT_NAVY,
        fg=WHITE,
        font=bold_font,
    )
    percent_label.pack()

    info_label = tk.Label(
        prog,
        text=f"0 of {total} | ETA: --:--",
        bg=GT_NAVY,
        fg=WHITE,
        font=_choose_font(prog),
    )
    info_label.pack(pady=(0, 20))

    prog.update()

    return prog, bar, percent_label, info_label


# ----------------------------
# Badge Generation Functions
# ----------------------------

def generate_labels(file_path: Path, save_path: Path, args, cfg: dict) -> None:
    """
    Generate QR badges on Avery 5395 (2×4 grid, 8 per page).
    Always saves DOCX and converts to PDF.
    """
    records = load_records(file_path)

    def safe_str(val):
        return str(val) if val is not None else ''

    records.sort(key=lambda r: (
        safe_str(r.get('Last', '')).lower(),
        safe_str(r.get('Preferred', '')).lower()
    ))

    total = len(records)
    start_time = time.time()

    doc = Document()
    sec = doc.sections[0]
    m = cfg['margins']
    sec.top_margin = Inches(m['top'])
    sec.bottom_margin = Inches(m['bottom'])
    sec.left_margin = Inches(m['left'])
    sec.right_margin = Inches(m['right'])
    if args.landscape:
        sec.orientation = WD_ORIENT.LANDSCAPE

    per_page, cols = cfg['per_page_qr'], cfg['columns_qr']
    rows = per_page // cols
    temp_files = []

    prog, bar, percent_label, info_label = create_progress_window("GT QR Badges Generator", total)

    try:
        for idx, rec in enumerate(records, start=1):
            if _tk_available and prog:
                elapsed = time.time() - start_time
                avg_per = elapsed / idx
                remaining = total - idx
                eta_seconds = int(avg_per * remaining)
                eta_str = (
                    time.strftime('%M:%S', time.gmtime(eta_seconds))
                    if eta_seconds < 3600
                    else time.strftime('%H:%M:%S', time.gmtime(eta_seconds))
                )

                bar['value'] = idx
                pct = int((idx / total) * 100)
                percent_label.config(text=f"{pct}%")
                info_label.config(text=f"{idx} of {total} | ETA: {eta_str}")
                prog.update()

            if (idx - 1) % per_page == 0:
                table = doc.add_table(rows=rows, cols=cols)
                table.autofit = False
                for col in table.columns:
                    for cell in col.cells:
                        cell.width = Inches(cfg['label_dims_in_qr']['width'])
                for row in table.rows:
                    row.height = Inches(cfg['label_dims_in_qr']['height'])
                    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

            r, c = divmod((idx - 1) % per_page, cols)
            cell = table.rows[r].cells[c]

            count_val = safe_str(rec.get('FASET Total Count', '')).strip()
            count_display = count_val if count_val else '1'

            p = cell.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            name_text = (
                f"{safe_str(rec.get('Preferred', '')).strip()} "
                f"{safe_str(rec.get('Last', '')).strip()} - {count_display}"
            ).strip()
            run = p.add_run(name_text)
            run.bold = True
            run.font.size = Pt(cfg['font_size_pt'])
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            shirt = safe_str(rec.get('FASET Shirt Size', '')).strip()
            if shirt:
                ps = cell.add_paragraph()
                ps.paragraph_format.space_before = Pt(0)
                ps.paragraph_format.space_after = Pt(0)
                rs = ps.add_run(f"Shirt Size: {shirt}")
                rs.font.size = Pt(cfg['font_size_pt'])
                ps.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            qr_url = safe_str(rec.get('Code', '')).strip()
            if qr_url:
                img_path = download_qr_image(qr_url, idx)
                if img_path:
                    temp_files.append(img_path)
                    ip = cell.add_paragraph()
                    ip.paragraph_format.space_before = Pt(0)
                    ip.paragraph_format.space_after = Pt(0)
                    ip.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    ip.add_run().add_picture(
                        str(img_path),
                        width=Inches(cfg['qr_size_in'])
                    )
                else:
                    err = cell.add_paragraph('QR download error')
                    err.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        tp = doc.add_paragraph()
        tp.paragraph_format.space_before = Pt(0)
        tp.paragraph_format.space_after = Pt(0)
        rt = tp.add_run(f"Total Count: {total}")
        rt.bold = True
        tp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        doc.save(save_path)

        for tmp in temp_files:
            try:
                tmp.unlink()
            except Exception:
                logging.warning(f"Could not delete temp file {tmp}")

    finally:
        if prog:
            prog.destroy()

    outdir = save_path.parent
    pdf_path = convert_to_pdf(save_path, outdir)
    open_file(save_path, background=False)
    if pdf_path:
        open_file(pdf_path, background=False)
    else:
        if _tk_available:
            messagebox.showerror('PDF Export Failed', 'Could not convert to PDF.')
        else:
            print('Warning: PDF conversion failed. DOCX saved at:', save_path)


def name_badges_fixed(file_path: Path, save_path: Path, cfg: dict = DEFAULT_CONFIG) -> None:
    """
    Generate Student Name badges on Avery 5392 (2×3 grid, 6 per page).
    Always saves DOCX and converts to PDF.
    """
    records = load_records(file_path)

    def safe_str(val):
        return str(val) if val is not None else ''

    records.sort(key=lambda r: (
        safe_str(r.get('Last', '')).lower(),
        safe_str(r.get('Preferred', '')).lower()
    ))

    total = len(records)
    start_time = time.time()

    doc = Document()
    sec = doc.sections[0]
    sec.page_height = Inches(11)
    sec.page_width = Inches(8.5)
    sec.top_margin = Inches(1.0)
    sec.bottom_margin = Inches(1.0)
    sec.left_margin = Inches(0.25)
    sec.right_margin = Inches(0.25)

    per_page = cfg['per_page_name']
    cols = cfg['columns_name']

    prog, bar, percent_label, info_label = create_progress_window("GT Student Name Badges Generator", total)

    try:
        for idx, rec in enumerate(records, start=1):
            if _tk_available and prog:
                elapsed = time.time() - start_time
                avg_per = elapsed / idx
                remaining = total - idx
                eta_seconds = int(avg_per * remaining)
                eta_str = (
                    time.strftime('%M:%S', time.gmtime(eta_seconds))
                    if eta_seconds < 3600
                    else time.strftime('%H:%M:%S', time.gmtime(eta_seconds))
                )

                bar['value'] = idx
                pct = int((idx / total) * 100)
                percent_label.config(text=f"{pct}%")
                info_label.config(text=f"{idx} of {total} | ETA: {eta_str}")
                prog.update()

            if (idx - 1) % per_page == 0:
                rows = per_page // cols
                table = doc.add_table(rows=rows, cols=cols)
                table.autofit = False
                for cell in table.columns[0].cells:
                    cell.width = Inches(cfg['label_dims_in_name']['width'])
                for cell in table.columns[1].cells:
                    cell.width = Inches(cfg['label_dims_in_name']['width'])
                for row in table.rows:
                    row.height = Inches(cfg['label_dims_in_name']['height'])
                    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

            r, c = divmod((idx - 1) % per_page, cols)
            cell = table.rows[r].cells[c]

            # Insert three blank paragraphs before content
            cell.add_paragraph()
            cell.add_paragraph()
            cell.add_paragraph()

            def add_centered_paragraph(text, font_size, bold=False):
                p = cell.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                run = p.add_run(text)
                run.bold = bold
                run.font.size = Pt(font_size)
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            name_text = (
                f"{safe_str(rec.get('Preferred', '')).strip()} "
                f"{safe_str(rec.get('Last', '')).strip()}"
            )
            if name_text.strip():
                add_centered_paragraph(name_text, font_size=15, bold=True)

            major = safe_str(rec.get('Major', '')).strip()
            if major:
                add_centered_paragraph(major, font_size=11)

            home_city = safe_str(rec.get('Home City', '')).strip()
            home_state = safe_str(rec.get('Home State/Region', '')).strip()
            if home_city or home_state:
                cs_text = f"{home_city}, {home_state}".strip(', ')
                add_centered_paragraph(cs_text, font_size=11)

            group = safe_str(rec.get('Group Number', '')).strip()
            pronouns = safe_str(rec.get('Pronouns', '')).strip()
            gp_text = ' – '.join(filter(None, [group, pronouns]))
            if gp_text:
                add_centered_paragraph(gp_text, font_size=11)

            session_date = safe_str(rec.get('FASET Session Date', '')).strip()
            if session_date:
                add_centered_paragraph(session_date, font_size=11)

        doc.save(save_path)

    finally:
        if prog:
            prog.destroy()

    outdir = save_path.parent
    pdf_path = convert_to_pdf(save_path, outdir)
    open_file(save_path, background=True)
    if pdf_path:
        open_file(pdf_path, background=False)
    else:
        if _tk_available:
            messagebox.showerror('PDF Export Failed', 'Could not convert to PDF.')
        else:
            print('Warning: PDF conversion failed. DOCX saved at:', save_path)


def guest1_badges(file_path: Path, save_path: Path, cfg: dict = DEFAULT_CONFIG) -> None:
    """
    Generate Guest 1 Name badges on Avery 5392 (2×3 grid, 6 per page).
    Uses fields:
      - Guest 1 Preferred Name, Guest 1 Last Name
      - Preferred, Last  (for “Guest of: Preferred Last”)
      - Guest 1 Affiliations
      - Home City, Home State/Region
      - FASET Session Date
    Skips any row where FASET Total Guest Count < 1.
    """
    records = load_records(file_path)

    def safe_str(val):
        return str(val) if val is not None else ''

    # Filter to those with at least 1 guest
    filtered = [
        r for r in records
        if safe_int(safe_str(r.get('FASET Total Guest Count', '0'))) >= 1
    ]

    # Sort by student Preferred+Last (same as student badge order)
    filtered.sort(key=lambda r: (
        safe_str(r.get('Last', '')).lower(),
        safe_str(r.get('Preferred', '')).lower()
    ))

    total = len(filtered)
    if total == 0:
        return

    start_time = time.time()

    doc = Document()
    sec = doc.sections[0]
    sec.page_height = Inches(11)
    sec.page_width = Inches(8.5)
    sec.top_margin = Inches(1.0)
    sec.bottom_margin = Inches(1.0)
    sec.left_margin = Inches(0.25)
    sec.right_margin = Inches(0.25)

    per_page = cfg['per_page_name']
    cols = cfg['columns_name']

    prog, bar, percent_label, info_label = create_progress_window("GT Guest 1 Name Badges Generator", total)

    try:
        for idx, rec in enumerate(filtered, start=1):
            if _tk_available and prog:
                elapsed = time.time() - start_time
                avg_per = elapsed / idx
                remaining = total - idx
                eta_seconds = int(avg_per * remaining)
                eta_str = (
                    time.strftime('%M:%S', time.gmtime(eta_seconds))
                    if eta_seconds < 3600
                    else time.strftime('%H:%M:%S', time.gmtime(eta_seconds))
                )

                bar['value'] = idx
                pct = int((idx / total) * 100)
                percent_label.config(text=f"{pct}%")
                info_label.config(text=f"{idx} of {total} | ETA: {eta_str}")
                prog.update()

            if (idx - 1) % per_page == 0:
                rows = per_page // cols
                table = doc.add_table(rows=rows, cols=cols)
                table.autofit = False
                for cell in table.columns[0].cells:
                    cell.width = Inches(cfg['label_dims_in_name']['width'])
                for cell in table.columns[1].cells:
                    cell.width = Inches(cfg['label_dims_in_name']['width'])
                for row in table.rows:
                    row.height = Inches(cfg['label_dims_in_name']['height'])
                    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

            r, c = divmod((idx - 1) % per_page, cols)
            cell = table.rows[r].cells[c]

            # Insert three blank paragraphs before content
            cell.add_paragraph()
            cell.add_paragraph()
            cell.add_paragraph()

            def add_centered_paragraph(text, font_size, bold=False):
                p = cell.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                run = p.add_run(text)
                run.bold = bold
                run.font.size = Pt(font_size)
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            guest1_name = (
                f"{safe_str(rec.get('Guest 1 Preferred Name', '')).strip()} "
                f"{safe_str(rec.get('Guest 1 Last Name', '')).strip()}"
            )
            if guest1_name.strip():
                add_centered_paragraph(guest1_name, font_size=15, bold=True)

            # “Guest of:” now shows Preferred + Last
            primary_preferred = safe_str(rec.get('Preferred', '')).strip()
            primary_last = safe_str(rec.get('Last', '')).strip()
            if primary_preferred or primary_last:
                add_centered_paragraph(f"Guest of: {primary_preferred} {primary_last}", font_size=11)

            guest1_aff = safe_str(rec.get('Guest 1 Affiliations', '')).strip()
            if guest1_aff:
                add_centered_paragraph(guest1_aff, font_size=11)

            home_city = safe_str(rec.get('Home City', '')).strip()
            home_state = safe_str(rec.get('Home State/Region', '')).strip()
            if home_city or home_state:
                cs_text = f"{home_city}, {home_state}".strip(', ')
                add_centered_paragraph(cs_text, font_size=11)

            session_date = safe_str(rec.get('FASET Session Date', '')).strip()
            if session_date:
                add_centered_paragraph(session_date, font_size=11)

        doc.save(save_path)
    finally:
        if prog:
            prog.destroy()

    outdir = save_path.parent
    pdf_path = convert_to_pdf(save_path, outdir)
    open_file(save_path, background=True)
    if pdf_path:
        open_file(pdf_path, background=False)
    else:
        if _tk_available:
            messagebox.showerror('PDF Export Failed', 'Could not convert to PDF.')
        else:
            print('Warning: PDF conversion failed. DOCX saved at:', save_path)


def guest2_badges(file_path: Path, save_path: Path, cfg: dict = DEFAULT_CONFIG) -> None:
    """
    Generate Guest 2 Name badges on Avery 5392 (2×3 grid, 6 per page).
    Uses fields:
      - Guest 2 Preferred Name, Guest 2 Last Name
      - Preferred, Last (for “Guest of: Preferred Last”)
      - Guest 2 Affiliations
      - Home City, Home State/Region
      - FASET Session Date
    Skips any row where FASET Total Guest Count < 2.
    """
    records = load_records(file_path)

    def safe_str(val):
        return str(val) if val is not None else ''

    filtered = [
        r for r in records
        if safe_int(safe_str(r.get('FASET Total Guest Count', '0'))) >= 2
    ]

    # Sort by student Preferred+Last (same as student badge order)
    filtered.sort(key=lambda r: (
        safe_str(r.get('Last', '')).lower(),
        safe_str(r.get('Preferred', '')).lower()
    ))

    total = len(filtered)
    if total == 0:
        return

    start_time = time.time()

    doc = Document()
    sec = doc.sections[0]
    sec.page_height = Inches(11)
    sec.page_width = Inches(8.5)
    sec.top_margin = Inches(1.0)
    sec.bottom_margin = Inches(1.0)
    sec.left_margin = Inches(0.25)
    sec.right_margin = Inches(0.25)

    per_page = cfg['per_page_name']
    cols = cfg['columns_name']

    prog, bar, percent_label, info_label = create_progress_window("GT Guest 2 Name Badges Generator", total)

    try:
        for idx, rec in enumerate(filtered, start=1):
            if _tk_available and prog:
                elapsed = time.time() - start_time
                avg_per = elapsed / idx
                remaining = total - idx
                eta_seconds = int(avg_per * remaining)
                eta_str = (
                    time.strftime('%M:%S', time.gmtime(eta_seconds))
                    if eta_seconds < 3600
                    else time.strftime('%H:%M:%S', time.gmtime(eta_seconds))
                )

                bar['value'] = idx
                pct = int((idx / total) * 100)
                percent_label.config(text=f"{pct}%")
                info_label.config(text=f"{idx} of {total} | ETA: {eta_str}")
                prog.update()

            if (idx - 1) % per_page == 0:
                rows = per_page // cols
                table = doc.add_table(rows=rows, cols=cols)
                table.autofit = False
                for cell in table.columns[0].cells:
                    cell.width = Inches(cfg['label_dims_in_name']['width'])
                for cell in table.columns[1].cells:
                    cell.width = Inches(cfg['label_dims_in_name']['width'])
                for row in table.rows:
                    row.height = Inches(cfg['label_dims_in_name']['height'])
                    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

            r, c = divmod((idx - 1) % per_page, cols)
            cell = table.rows[r].cells[c]

            cell.add_paragraph()
            cell.add_paragraph()
            cell.add_paragraph()

            def add_centered_paragraph(text, font_size, bold=False):
                p = cell.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                run = p.add_run(text)
                run.bold = bold
                run.font.size = Pt(font_size)
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            guest2_name = (
                f"{safe_str(rec.get('Guest 2 Preferred Name', '')).strip()} "
                f"{safe_str(rec.get('Guest 2 Last Name', '')).strip()}"
            )
            if guest2_name.strip():
                add_centered_paragraph(guest2_name, font_size=15, bold=True)

            # “Guest of:” now shows Preferred + Last
            primary_preferred = safe_str(rec.get('Preferred', '')).strip()
            primary_last = safe_str(rec.get('Last', '')).strip()
            if primary_preferred or primary_last:
                add_centered_paragraph(f"Guest of: {primary_preferred} {primary_last}", font_size=11)

            guest2_aff = safe_str(rec.get('Guest 2 Affiliations', '')).strip()
            if guest2_aff:
                add_centered_paragraph(guest2_aff, font_size=11)

            home_city = safe_str(rec.get('Home City', '')).strip()
            home_state = safe_str(rec.get('Home State/Region', '')).strip()
            if home_city or home_state:
                cs_text = f"{home_city}, {home_state}".strip(', ')
                add_centered_paragraph(cs_text, font_size=11)

            session_date = safe_str(rec.get('FASET Session Date', '')).strip()
            if session_date:
                add_centered_paragraph(session_date, font_size=11)

        doc.save(save_path)
    finally:
        if prog:
            prog.destroy()

    outdir = save_path.parent
    pdf_path = convert_to_pdf(save_path, outdir)
    open_file(save_path, background=True)
    if pdf_path:
        open_file(pdf_path, background=False)
    else:
        if _tk_available:
            messagebox.showerror('PDF Export Failed', 'Could not convert to PDF.')
        else:
            print('Warning: PDF conversion failed. DOCX saved at:', save_path)


# ----------------------------
# Main
# ----------------------------

def main() -> None:
    try:
        args = parse_args()

        config = DEFAULT_CONFIG.copy()
        if args.config:
            loaded = load_config(Path(args.config))
            if loaded:
                config.update(loaded)

        if _tk_available:
            root = Tk()
            root.title('GT QR & Name Badge Generator')
            root.configure(bg=GT_NAVY)
            root.resizable(False, False)

            font_family, font_size = _choose_font(root)
            large_font = (font_family, 14, 'bold')

            style = Style(root)
            style.theme_use('clam')
            style.configure(
                'Flat.TButton',
                background=GT_GOLD,
                foreground=GT_NAVY,
                font=large_font,
                relief='flat',
                borderwidth=0
            )
            style.map(
                'Flat.TButton',
                background=[('active', GT_HOVER)],
                foreground=[('active', GT_NAVY)]
            )

            # ─── Full GT logo at top of dialog ─────────────────────────
            img_full = _load_image_file(FULL_LOGO_FILE) or _load_embedded_image(FULL_LOGO_B64)
            if img_full:
                # Resize to ~300px wide, preserving aspect ratio
                w_percent = (300 / float(img_full.size[0]))
                hsize = int((float(img_full.size[1]) * float(w_percent)))
                logo_img = img_full.resize((300, hsize), RESAMPLE)
                logo_tk = ImageTk.PhotoImage(logo_img)
                logo_lbl = tk.Label(root, image=logo_tk, bg=GT_NAVY)
                logo_lbl.image = logo_tk
                logo_lbl.pack(pady=(20, 10))
            else:
                # Fallback text if decoding fails or PIL not available
                fallback_lbl = tk.Label(
                    root,
                    text='Georgia Tech',
                    bg=GT_NAVY,
                    fg=GT_GOLD,
                    font=(font_family, 24, 'bold')
                )
                fallback_lbl.pack(pady=(20, 10))

            # ─── Ribbon/pattern under the logo ───────────────────────
            img_ribbon = _load_image_file(RIBBON_FILE) or _load_embedded_image(RIBBON_B64)
            if img_ribbon:
                # Stretch ribbon to ~half-screen width
                dialog_width = root.winfo_screenwidth() // 2
                w_percent = (dialog_width / float(img_ribbon.size[0]))
                hsize = int((float(img_ribbon.size[1]) * float(w_percent)))
                ribbon_img = img_ribbon.resize((dialog_width, hsize), RESAMPLE)
                ribbon_tk = ImageTk.PhotoImage(ribbon_img)
                ribbon_lbl = tk.Label(root, image=ribbon_tk, bg=GT_NAVY)
                ribbon_lbl.image = ribbon_tk
                ribbon_lbl.pack(pady=(0, 20))
            else:
                fallback_ribbon = tk.Label(
                    root,
                    text='​[Ribbon Missing]​',
                    bg=GT_NAVY,
                    fg=WHITE,
                    font=_choose_font(root)
                )
                fallback_ribbon.pack(pady=(0, 20))

            template_var = StringVar(value='')

            def select_and_close(value):
                template_var.set(value)
                root.destroy()

            btn_qr = TtkButton(
                root,
                text='QR Badges Only',
                command=lambda: select_and_close('qr'),
                style='Flat.TButton',
                padding=(20, 12)
            )
            btn_qr.pack(fill='x', padx=60, pady=12)

            btn_student = TtkButton(
                root,
                text='Student Name Badges Only',
                command=lambda: select_and_close('student'),
                style='Flat.TButton',
                padding=(20, 12)
            )
            btn_student.pack(fill='x', padx=60, pady=12)

            btn_guest1 = TtkButton(
                root,
                text='Guest 1 Name Badges Only',
                command=lambda: select_and_close('guest1'),
                style='Flat.TButton',
                padding=(20, 12)
            )
            btn_guest1.pack(fill='x', padx=60, pady=12)

            btn_guest2 = TtkButton(
                root,
                text='Guest 2 Name Badges Only',
                command=lambda: select_and_close('guest2'),
                style='Flat.TButton',
                padding=(20, 12)
            )
            btn_guest2.pack(fill='x', padx=60, pady=12)

            btn_all = TtkButton(
                root,
                text='Print All Badges',
                command=lambda: select_and_close('all'),
                style='Flat.TButton',
                padding=(20, 12)
            )
            btn_all.pack(fill='x', padx=60, pady=(12, 60))

            # Center window on screen and enlarge
            root.update_idletasks()
            w = root.winfo_width()
            h = root.winfo_height()
            w_enlarged = int(w * 1.2)
            h_enlarged = int(h * 1.2)
            ws = root.winfo_screenwidth()
            hs = root.winfo_screenheight()
            x = (ws // 2) - (w_enlarged // 2)
            y = (hs // 2) - (h_enlarged // 2)
            root.geometry(f'{w_enlarged}x{h_enlarged}+{x}+{y}')

            root.mainloop()
            template = template_var.get()

            if not template:
                sys.exit(0)
        else:
            template = 'qr'

        file_path = Path(args.input) if args.input else pick_input_file()
        records = load_records(file_path)
        raw_date = str(records[0].get('FASET Session Date', '')).strip()
        if raw_date:
            sanitized_date = raw_date.replace(' ', '_').replace('/', '-')
        else:
            sanitized_date = file_path.stem

        if template == 'qr':
            default_name = f"{sanitized_date}_QRBadges.docx"
            save_path = Path(args.output) if args.output else pick_save_file(default_name)
            generate_labels(file_path, save_path, args, config)

        elif template == 'student':
            default_name = f"{sanitized_date}_StudentNameBadges.docx"
            save_path = Path(args.output) if args.output else pick_save_file(default_name)
            name_badges_fixed(file_path, save_path, config)

        elif template == 'guest1':
            default_name = f"{sanitized_date}_Guest1Badge.docx"
            save_path = Path(args.output) if args.output else pick_save_file(default_name)
            guest1_badges(file_path, save_path, config)

        elif template == 'guest2':
            default_name = f"{sanitized_date}_Guest2Badge.docx"
            save_path = Path(args.output) if args.output else pick_save_file(default_name)
            guest2_badges(file_path, save_path, config)

        else:  # 'all'
            qr_name = f"{sanitized_date}_QRBadges.docx"
            qr_save = pick_save_file(qr_name)
            generate_labels(file_path, qr_save, args, config)

            student_save = qr_save.parent / f"{sanitized_date}_StudentNameBadges.docx"
            name_badges_fixed(file_path, student_save, config)

            guest1_save = qr_save.parent / f"{sanitized_date}_Guest1Badge.docx"
            guest1_badges(file_path, guest1_save, config)

            guest2_save = qr_save.parent / f"{sanitized_date}_Guest2Badge.docx"
            guest2_badges(file_path, guest2_save, config)

    except Exception:
        logging.exception('Unhandled exception in main')
        if _tk_available:
            try:
                messagebox.showerror('Error', f"An error occurred:\n{traceback.format_exc()}")
            except:
                pass
        else:
            print('Error:', traceback.format_exc(), file=sys.stderr)
        sys.exit(1)


if __name__ == '__main__':
    main()

# ----------------------------
# Simple Test Case (bottom of script)
# ----------------------------
# To test, create a dummy CSV named 'test_dummy.csv' containing:
# Preferred,Last,FASET Total Guest Count,FASET Shirt Size,Code,Major,Home City,Home State/Region,Group Number,Pronouns,FASET Session Date,Guest 1 Preferred Name,Guest 1 Last Name,Guest 1 Affiliations,Guest 2 Preferred Name,Guest 2 Last Name,Guest 2 Affiliations
# John,Doe,2,M,http://example.com/qr1.png,CS,Atlanta,GA,1,he/him,2025-08-01,Jane,Doe,Friend,Mark,Doe,Colleague
#
# Then run:
#     python3 generate_qr_name_badges.py -i test_dummy.csv
